

import openpyxl
from docx import Document
import re


BILL = r"(Рахунок\s+?№\s+?)([_]+)\b"
DATE = r"(Дата:?\s+?)(__.__.____)\b"
BUYER = r"(Покупець:?\s+?)([_]+)\b"
TOTAL = r"(Всього:?\s+?)([_]+)\b"


def create_bill(bill_no, data, template):
    wb = openpyxl.load_workbook(data)

    ws = wb["bills"]
    for row in ws.rows:
        if row[1].value == str(bill_no):
            bill_id, no, date, client_id = [c.value for c in row]
            # print(bill_id, no, date, client_id)
            break

    ws = wb["buyers"]
    for row in ws.rows:
        if row[0].value == client_id:
            client_id, client_name, address = [c.value for c in row]
            # print(client_id, client_name, address)
            break

    products = {}
    ws = wb["items"]
    for row in ws.rows:
        if row[0].value == bill_id:
            bill_id, product_id, quantity = [c.value for c in row]
            # print(bill_id, product_id, quantity)
            products[product_id] = {"quantity": quantity}

    # print(products)

    ws = wb["products"]
    for row in ws.rows:
        product_id = row[0].value
        if product_id in products:
            product_id, name, unit, price = [c.value for c in row]
            # print(product_id, name, unit, price)
            products[product_id]["name"] = name
            products[product_id]["unit"] = unit
            products[product_id]["price"] = price

    print(products)

    doc = Document(template)

    total = 0
    t = doc.tables[0]
    for i, product in enumerate(products, 1):
        full_price = (float(products[product]["quantity"]) *
                      float(products[product]["price"]))
        total += full_price
        values = (str(i),
                  products[product]["name"],
                  products[product]["unit"],
                  products[product]["quantity"],
                  products[product]["price"],
                  str(full_price))
        row = t.add_row()
        for cell, value in zip(row.cells, values):
            cell.text = value

    def change_bill_no(match):
        count = match.group(2).count("_")
        s = "{:_>" + str(count) + "}"
        return match.group(1) + s.format(bill_no)

    def change_date(match):
        return match.group(1) + str(date)

    def change_buyer(match):
        return match.group(1) + client_name + " " + address

    def change_total(match):
        return match.group(1) + str(total)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.text = re.sub(BILL, change_bill_no, run.text)
            run.text = re.sub(DATE, change_date, run.text)
            run.text = re.sub(BUYER, change_buyer, run.text)
            run.text = re.sub(TOTAL, change_total, run.text)

    doc.save("bill_" + client_name + "_" + date + ".docx")


if __name__ == '__main__':
    create_bill(253, "data.xlsx", "template.docx")
