import openpyxl
from docx import Document
from docx.shared import Pt, Cm


def set_column_width(column, width):
    """ Встановлює ширину рядка таблиці"""
    for cell in column.cells:
        cell.width = width


def create_invoice(invoice_no: int, datafile: str):
    """ Створює рахунок за номером"""
    wb = openpyxl.load_workbook(datafile)

    ws = wb["invoices"]
    for row in ws.rows:
        if row[1].value == invoice_no:
            invoice_id, no, date, customer_id = [c.value for c in row]
            # print(invoice_id, no, date, customer_id)
            break
    else:
        return

    ws = wb["customers"]
    for row in ws.rows:
        if row[0].value == customer_id:
            customer_id, customer_name, address = [c.value for c in row]
            # print(customer_id, customer_name, address)
            break
    else:
        return

    products = {}
    ws = wb["items"]
    for row in ws.rows:
        if row[0].value == invoice_id:
            invoice_id, product_id, quantity = [c.value for c in row]
            # print(invoice_id, product_id, quantity)
            products[product_id] = {"quantity": quantity}

    # print(products)

    ws = wb["products"]
    for row in ws.rows:
        product_id = row[0].value
        if product_id in products:
            product_id, name, unit, price = [c.value for c in row]
            # print(product_id, name, unit, price)
            products[product_id]["name"] = name
            products[product_id]["unit"] = unit
            products[product_id]["price"] = price

    # print(products)

    # Створюємо документ Word
    doc = Document()
    # Задаємо загальний стиль документа
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"  # Тип шрифту
    style.font.size = Pt(14)  # Розмір шрифту

    # Додаємо номер рахунку та дату
    table = doc.add_table(1, 2)
    row = table.rows[0]
    paragraph = row.cells[0].add_paragraph(f"Рахунок № {invoice_no}")
    paragraph.alignment = 0
    paragraph = row.cells[1].add_paragraph(f"Дата {date}")
    paragraph.alignment = 2

    # Додаємо ім`я покупця
    doc.add_paragraph(f"Покупець: {customer_name} {address}")

    headers = ("№", "Назва", "Од. виміру", "Кількість", "Ціна", "Сума")
    # Створюємо таблицю
    table = doc.add_table(1, len(headers), "Table Grid")
    # Додаємо заголовок
    row = table.rows[0]
    for cell, header in zip(row.cells, headers):
        cell.text = header
    # Додаємо рядки до таблиці та знаходимо загальну вартість
    total = 0
    for i, product in enumerate(products.values(), 1):
        full_price = float(product["quantity"]) * float(product["price"])
        # Знаходимо загальну вартість
        total += full_price
        values = (
            i,
            product["name"],
            product["unit"],
            product["quantity"],
            product["price"],
            full_price
        )
        # Додаємо рядки до таблиці
        row = table.add_row()
        for cell, value in zip(row.cells, values):
            cell.text = str(value)

    # Всановлюємо розміри клітинок таблиці
    widths = (Cm(2.0), Cm(5.0), Cm(3.0), Cm(1.5), Cm(1.5))
    for column, width in zip(table.columns, widths):
        set_column_width(column, width)

    doc.add_paragraph()
    # Додаємо загальну вартість
    paragraph = doc.add_paragraph(f"Всього: {total}")
    paragraph.alignment = 2

    doc.save(f"invoice_{customer_name}_{''.join(date.split('.'))}.docx")


if __name__ == "__main__":
    create_invoice(253, "data.xlsx")
